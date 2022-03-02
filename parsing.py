import re
import os.path
from os import path
from docx import Document
from docx.shared import Inches
from PIL import ImageTk, Image
from tkinter import Tk, filedialog, Frame, Button, Label, Entry, ttk, PhotoImage, LabelFrame, Y, RIGHT, Scrollbar, NS, DISABLED, NORMAL

class Gui:
    def __init__(self, master):
        self.master = master
        self.master.minsize(width=650, height=250)
        self.file_path = None
        self.wordHandle = None
        self.input_frame = None
        self.file_frame = None
        self.button_frame = None
        self.create_widgets()
        self.inputObject = {
            'text':{},
            'image': {}
        }

    def create_widgets(self):
        # Input frame
        self.input_frame = self.create_input_frame()
        self.input_frame.grid(column=0, row=1, padx=8, pady=5, sticky="nsew")
        
        # File frame
        self.file_frame = self.create_file_frame()
        self.file_frame.grid(column=0, row=0, padx=8, pady=5, sticky="nsew", columnspan=2, ipady=5)

        # Button frame
        self.button_frame = self.create_button_frame()
        self.button_frame.grid(column=1, row=1, padx=8, pady=5, sticky="nsew")

    def create_file_frame(self):
        frame = ttk.LabelFrame(self.master, text="Target File")
        # grid layout for the input frame
        frame.columnconfigure(0, weight=1)
        frame.columnconfigure(1, weight=2)
        frame.columnconfigure(2, weight=1)
        
        Label(frame, text="File path ").grid(row=0)
        Button(frame, text="Select", command=self.select_file).grid(row=0, column=2)
        for widget in frame.winfo_children():
            widget.grid(ipadx=10)
        
        return frame

    def create_input_frame(self):
        frame = ttk.LabelFrame(self.master, text="Replacement")
        # grid layout for the input frame
        frame.columnconfigure(0, weight=1)
        frame.columnconfigure(1, weight=1)
        frame.columnconfigure(2, weight=2)

        for widget in frame.winfo_children():
            widget.grid(padx=3, pady=5)
        
        return frame

    def create_button_frame(self):
        frame = ttk.LabelFrame(self.master, text="Target File")
        frame.columnconfigure(0, weight=1)
        self.saveAsBtn = Button(frame, text="Save", command=self.file_save, state=DISABLED)
        self.saveAsBtn.grid(row=0, column=0)
        Button(frame, text="Quit", command=self.master.destroy).grid(row=1, column=0)

        for widget in frame.winfo_children():
            widget.grid(padx=0, pady=10)
        
        return frame

    def add_entry(self, frame):
        row = 0
        for tag_type, tags in self.wordHandle.keyword.items():
            for tag in list(tags.keys()):
                Label(frame, text=tag+": ").grid(row=row)
                if tag_type == 'text':
                    self.inputObject[tag_type][tag] = Entry(frame)
                    self.inputObject[tag_type][tag].grid(row=row, column=1)
                elif tag_type == 'image':
                    Button(frame, text="Choose", command=lambda row=row, tag=tag:self.select_image(row, tag)).grid(row=row, column=1)            
                row+=1
        
        for widget in frame.winfo_children():
            widget.grid(padx=0, pady=6)

    def select_file(self): 
        file_path = filedialog.askopenfilename(title='Select file',initialdir='~/', filetypes = (("docx files","*.docx"),("all files","*.*")))
        self.wordHandle = WordHandle(file_path)
        Label(self.file_frame, text=file_path).grid(row=0, column=1, ipadx=10)
        self.add_entry(self.input_frame)
        self.saveAsBtn.config(state=NORMAL)

    def select_image(self, row, tag): 
        image_path = filedialog.askopenfilename(title='Choose',initialdir='~/', filetypes = (("jpeg files","*.jpg"),("png files","*.png"),("all files","*.*")))
        img = ImageTk.PhotoImage(Image.open(image_path).resize((100, 100), Image.ANTIALIAS))  # PIL solution
        cercaImg = Label(self.input_frame, image = img)
        cercaImg.grid(row=row, column=2, padx=3, pady=6)
        cercaImg.photo = img
        self.inputObject['image'][tag] = image_path

    def file_save(self):
        f = filedialog.asksaveasfile(mode='w', defaultextension=".docx", filetypes=(("Word", "*.docx"), ("All files", "*")))
        if f is None:
            return
        for label, entry in self.inputObject['text'].items():
            self.inputObject['text'][label] = entry.get()
        
        self.wordHandle.docx_replace_regex(None, self.inputObject)
        self.wordHandle.doc.save(f.name)

class WordHandle:
    def __init__(self, filename):
        self.doc = Document(filename)
        self.keyword = {
            'text': self.getKeywordDict(self.doc, re.compile(r'\<.*?\>'), {}),
            'image': self.getKeywordDict(self.doc, re.compile(r'\[.*?\]'), {})
        }

    def getKeywordDict(self, doc_obj, regex, keyword):
        for p in doc_obj.paragraphs:
            if regex.search(p.text):
                inline = p.runs
                for i in range(len(inline)):
                    if regex.search(inline[i].text):
                        if regex.search(inline[i].text).group() not in keyword:
                            keyword[regex.search(inline[i].text).group()] = ''
        for table in doc_obj.tables:
            for row in table.rows:
                for cell in row.cells:
                    self.getKeywordDict(cell, regex, keyword)
        return keyword

    def docx_replace_regex(self, doc_obj=None, keywordDicts=None):
        if not doc_obj:
            doc_obj = self.doc
        if not keywordDicts:
            keywordDicts = self.keyword
        for p in doc_obj.paragraphs:
            for replace_type, keywordDict in keywordDicts.items():
                for keyword, replace in keywordDict.items():
                    if keyword in p.text:
                        inline = p.runs
                        for i in range(len(inline)):
                            if keyword in inline[i].text:
                                if replace_type == 'text':
                                    inline[i].text = inline[i].text.replace(keyword, replace)
                                elif replace_type == 'image':
                                    inline[i].text = ''
                                    self.insertImage(p, replace)
                                    
        for table in doc_obj.tables:
            for row in table.rows:
                for cell in row.cells:
                    self.docx_replace_regex(cell, keywordDicts)

    def insertImage(self, paragraph, file_path):
        r = paragraph.add_run()
        if not file_path:
            return
        r.add_picture(file_path, width=Inches(2.0), height=Inches(2.0))
    
def main():
    root = Tk()
    root.title('Parsing')
    root.resizable(0, 0)
    # windows only (remove the minimize/maximize button)
    root.attributes('-toolwindow', True)

    root.columnconfigure(0, weight=2)
    root.columnconfigure(1, weight=1)
    my_gui = Gui(root)
    root.mainloop()

if __name__== "__main__":
	main()

